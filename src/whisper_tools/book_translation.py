"""
Book Translation Module (Arabic → German / English / Both)
==========================================================

Supports translating book files (.txt, .pdf, .docx) using the 3-Tier
fallback translation pipeline with batching, live JSON backup, and resume.
Outputs:
  - Translated Text / Markdown book files
  - Word (.docx) documents (if python-docx is installed)
  - Full structured JSON backup
"""

import logging
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


def extract_text_from_book(file_path: Path) -> str:
    """Extract full raw text from PDF, DOCX, or TXT book file."""
    ext = file_path.suffix.lower()
    
    if ext == ".txt" or ext == ".md":
        return file_path.read_text(encoding="utf-8")
        
    elif ext == ".pdf":
        try:
            import pypdf
            reader = pypdf.PdfReader(str(file_path))
            text_parts = []
            for i, page in enumerate(reader.pages):
                p_text = page.extract_text()
                if p_text:
                    text_parts.append(p_text)
            return "\n\n".join(text_parts)
        except ImportError:
            raise RuntimeError("pypdf is required to read PDF books. Install with: uv add pypdf")
            
    elif ext == ".docx":
        try:
            import docx
            doc = docx.Document(str(file_path))
            return "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            raise RuntimeError("python-docx is required to read DOCX books. Install with: uv add python-docx")
            
    else:
        raise ValueError(f"Unsupported book format: {ext}. Supported formats: .txt, .pdf, .docx, .md")


def chunk_text_into_paragraphs(text: str, max_chunk_chars: int = 1000) -> list[dict[str, Any]]:
    """Split book text into coherent paragraph segments for translation."""
    raw_paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    
    chunks = []
    current_chunk = []
    current_length = 0
    
    for p in raw_paragraphs:
        if current_length + len(p) > max_chunk_chars and current_chunk:
            combined = "\n".join(current_chunk)
            chunks.append({
                "id": len(chunks) + 1,
                "text": combined,
                "original_ar": combined,
            })
            current_chunk = [p]
            current_length = len(p)
        else:
            current_chunk.append(p)
            current_length += len(p)
            
    if current_chunk:
        combined = "\n".join(current_chunk)
        chunks.append({
            "id": len(chunks) + 1,
            "text": combined,
            "original_ar": combined,
        })
        
    return chunks


def translate_book_interactive(book_path: Path, output_dir: Path, translate_segments_fn, args) -> Path:
    """Interactively process and translate a book from any source language to multiple target languages with dedicated directory & live backup."""
    base_name = book_path.stem
    book_output_dir = output_dir / "books" / base_name
    book_output_dir.mkdir(parents=True, exist_ok=True)
    
    # Determine Source & Target Languages dynamically
    if getattr(args, "source_lang_explicit", False) or getattr(args, "target_lang_explicit", False):
        source_lang = getattr(args, "source_lang", "ar").lower()
        target_langs_str = getattr(args, "target_lang", "de,en").lower()
        target_langs = [l.strip() for l in target_langs_str.split(",") if l.strip()]
    else:
        # Prompt Source Language
        print("\n  Select Book Source Language:")
        print("    [1] Arabic (ar - default)")
        print("    [2] German (de)")
        print("    [3] English (en)")
        print("    [4] Custom source language code (e.g. fr, es, tr)")
        src_choice = input("  Enter choice [1-4] (default: 1 Arabic): ").strip()
        if src_choice == "2":
            source_lang = "de"
        elif src_choice == "3":
            source_lang = "en"
        elif src_choice == "4":
            custom_src = input("  Enter custom source language code (e.g. fr, es): ").strip().lower()
            source_lang = custom_src if custom_src else "ar"
        else:
            source_lang = "ar"
            
        args.source_lang = source_lang

        # Prompt Target Languages based on Source
        print(f"\n  Select Target Translation Language(s) (Source: {source_lang.upper()}):")
        if source_lang == "ar":
            print("    [1] German (de)")
            print("    [2] English (en)")
            print("    [3] Both German & English (de,en - default)")
            print("    [4] Custom target language codes (comma-separated, e.g. de,en,fr)")
            lang_choice = input("  Enter choice [1-4] (default: 3 Both): ").strip()
            if lang_choice == "1":
                target_langs = ["de"]
            elif lang_choice == "2":
                target_langs = ["en"]
            elif lang_choice == "4":
                custom_in = input("  Enter comma-separated target language codes: ").strip().lower()
                target_langs = [l.strip() for l in custom_in.split(",") if l.strip()] if custom_in else ["de", "en"]
            else:
                target_langs = ["de", "en"]
        elif source_lang == "de":
            print("    [1] Arabic (ar - default)")
            print("    [2] English (en)")
            print("    [3] Both Arabic & English (ar,en)")
            print("    [4] Custom target language codes (comma-separated, e.g. ar,en,fr)")
            lang_choice = input("  Enter choice [1-4] (default: 1 Arabic): ").strip()
            if lang_choice == "2":
                target_langs = ["en"]
            elif lang_choice == "3":
                target_langs = ["ar", "en"]
            elif lang_choice == "4":
                custom_in = input("  Enter comma-separated target language codes: ").strip().lower()
                target_langs = [l.strip() for l in custom_in.split(",") if l.strip()] if custom_in else ["ar"]
            else:
                target_langs = ["ar"]
        elif source_lang == "en":
            print("    [1] Arabic (ar - default)")
            print("    [2] German (de)")
            print("    [3] Both Arabic & German (ar,de)")
            print("    [4] Custom target language codes (comma-separated, e.g. ar,de,fr)")
            lang_choice = input("  Enter choice [1-4] (default: 1 Arabic): ").strip()
            if lang_choice == "2":
                target_langs = ["de"]
            elif lang_choice == "3":
                target_langs = ["ar", "de"]
            elif lang_choice == "4":
                custom_in = input("  Enter comma-separated target language codes: ").strip().lower()
                target_langs = [l.strip() for l in custom_in.split(",") if l.strip()] if custom_in else ["ar"]
            else:
                target_langs = ["ar"]
        else:
            custom_in = input(f"  Enter comma-separated target language codes for {source_lang.upper()} source (default: en): ").strip().lower()
            target_langs = [l.strip() for l in custom_in.split(",") if l.strip()] if custom_in else ["en"]

        args.target_lang = ",".join(target_langs)
            
    print(f"\n{'='*60}")
    print("  🚀 STARTING BOOK TRANSLATION PIPELINE")
    print(f"{'='*60}")
    print(f"  📖 Book File Name    : {book_path.name}")
    print(f"  🌐 Source Language   : {source_lang.upper()}")
    print(f"  🎯 Target Languages  : {', '.join(t.upper() for t in target_langs)}")
    print(f"  📁 Output Directory  : {book_output_dir}")
    print(f"{'='*60}\n")
    
    # Check for existing live backup / progress
    progress_file = book_output_dir / "translation_progress.json"
    temp_json = book_output_dir / "translation_temp.json"
    if progress_file.exists():
        try:
            import json
            prog_data = json.loads(progress_file.read_text(encoding="utf-8"))
            if not isinstance(prog_data, dict):
                raise ValueError(f"Unexpected progress file format: {type(prog_data).__name__}")
            completed_cnt = prog_data.get("completed_count", 0)
            print(f"  [FOUND PREVIOUS BACKUP] Found translation progress for '{base_name}' ({completed_cnt} segments completed).")
            res_ans = input("  Resume from existing backup to speed up execution? (y/n, default: y): ").strip().lower()
            if res_ans == "n":
                progress_file.unlink(missing_ok=True)
                if temp_json.exists():
                    temp_json.unlink(missing_ok=True)
                print("  [INFO] Cleared previous backup. Starting fresh translation from scratch!\n")
            else:
                print("  [INFO] Resuming from existing backup to speed up translation!\n")
        except ValueError as exc:
            # Progress file has an unrecognised schema (e.g. old list format).
            # Delete it so it doesn't keep triggering this warning every run.
            logger.warning("Discarding malformed progress file (%s) — starting fresh.", exc)
            progress_file.unlink(missing_ok=True)
            if temp_json.exists():
                temp_json.unlink(missing_ok=True)
        except Exception as exc:
            logger.debug("Could not read resume progress file: %s", exc)
    
    # 2. Extract & Chunk Text
    print(f"\n[1/3] Extracting text from {book_path.name}...")
    raw_text = extract_text_from_book(book_path)
    
    # 2.5 Optional AI Verification / Proofreading of Extracted Text
    # The full book can easily exceed the LLM's context window (131k tokens),
    # so we split into small chunks and proofread each one independently.
    if getattr(args, "verify_book", True) and getattr(args, "local_translate", True):
        try:
            from whisper_tools.translation import verify_transcription_with_llm
        except ImportError:
            from .translation import verify_transcription_with_llm

        _PROOFREAD_CHUNK_CHARS = 2000  # ~500 tokens — well within model limits
        print("       [AI Verification] Proofreading extracted book text for OCR/formatting fixes...")
        try:
            chunks = [
                raw_text[i : i + _PROOFREAD_CHUNK_CHARS]
                for i in range(0, len(raw_text), _PROOFREAD_CHUNK_CHARS)
            ]
            proofread_chunks = []
            for idx, chunk in enumerate(chunks, 1):
                result = verify_transcription_with_llm(chunk, source_lang=source_lang)
                proofread_chunks.append(result if result and result.strip() else chunk)
                if idx % 10 == 0:
                    print(f"       [AI Verification] Proofread {idx}/{len(chunks)} chunks...")
            raw_text = "".join(proofread_chunks)
            print("       [AI Verification] Book text proofreading complete!")
        except Exception as e:
            print(f"       [NOTE] AI book text proofreading skipped: {e}")

    segments = chunk_text_into_paragraphs(raw_text)
    print(f"       Extracted {len(segments)} paragraph chunks from book.")
    
    # 3. Translate Segments via Pipeline with dedicated directory for live JSON backups
    print("\n[2/3] Translating book paragraphs using 3-Tier AI Engine...")
    translated_segments = translate_segments_fn(segments, book_output_dir, args, source_lang=source_lang, target_langs=target_langs)
    
    # 3.5 Full-Book AI Editorial Audit & Refinement Pass
    audit_reports = []
    if getattr(args, "local_translate", True):
        try:
            try:
                from whisper_tools.book_editorial import (
                    audit_and_refine_full_book,
                    write_editorial_audit_report,
                )
            except ImportError:
                from .book_editorial import (
                    audit_and_refine_full_book,
                    write_editorial_audit_report,
                )

            llm_model_id = getattr(args, "mlx_model", "mlx-community/Qwen2.5-7B-Instruct-4bit")
            for target in target_langs:
                translated_segments, audit_entries = audit_and_refine_full_book(
                    raw_text, translated_segments, source_lang, target, llm_model_id
                )
                rep_path = write_editorial_audit_report(audit_entries, book_output_dir, base_name, target, len(translated_segments))
                audit_reports.append(f"Editorial Audit Report ({target.upper()}): [`{rep_path.name}`](file://{rep_path.absolute()})")
        except Exception as e:
            print(f"       [NOTE] AI Editorial Audit pass skipped: {e}")

    # 4. Save Extracted Raw Text Transcript File
    extracted_src_path = book_output_dir / f"{base_name}_extracted_{source_lang}.txt"
    extracted_src_path.write_text(raw_text, encoding="utf-8")
    print(f"       Extracted {source_lang.upper()} Transcript (.txt): {extracted_src_path.name}")
 
    # 5. Save Output Book Files
    print("\n[3/3] Saving translated book deliverables & summary report...")
    
    # Save each target language txt/docx
    # Import e-book exporters
    from whisper_tools.book_exporter import enrich_markdown_arabic, export_to_epub, export_to_pdf

    # Save each target language txt/docx/pdf/epub
    generated_files = []
    for target in target_langs:
        t_text = "\n\n".join([s.get(f"text_{target}", s.get("text", "")) for s in translated_segments])
        
        # 1. Plain Text (.txt)
        txt_path = book_output_dir / f"{base_name}_translated_{target}.txt"
        txt_path.write_text(t_text, encoding="utf-8")
        print(f"       {target.upper()} Book (.txt): {txt_path.name}")
        generated_files.append(f"{target.upper()} Book (.txt): [`{txt_path.name}`](file://{txt_path.absolute()})")
        
        # 2. Word Document (.docx)
        try:
            import docx
            doc = docx.Document()
            doc.add_heading(f"{base_name} - {target.upper()} Translation", 0)
            for s in translated_segments:
                p_text = s.get(f"text_{target}", s.get("text", "")).strip()
                if p_text:
                    doc.add_paragraph(p_text)
            docx_path = book_output_dir / f"{base_name}_translated_{target}.docx"
            doc.save(str(docx_path))
            print(f"       {target.upper()} Book (.docx): {docx_path.name}")
            generated_files.append(f"{target.upper()} Book (.docx): [`{docx_path.name}`](file://{docx_path.absolute()})")
        except Exception as exc:
            logger.debug("DOCX export skipped: %s", exc)

        # 3. PDF E-Book (.pdf)
        pdf_path = book_output_dir / f"{base_name}_translated_{target}.pdf"
        if export_to_pdf(t_text, pdf_path, title=f"{base_name} ({target.upper()})"):
            generated_files.append(f"{target.upper()} Book (.pdf): [`{pdf_path.name}`](file://{pdf_path.absolute()})")

        # 4. EPUB E-Book (.epub)
        epub_path = book_output_dir / f"{base_name}_translated_{target}.epub"
        if export_to_epub(t_text, epub_path, title=f"{base_name} ({target.upper()})", language=target):
            generated_files.append(f"{target.upper()} Book (.epub): [`{epub_path.name}`](file://{epub_path.absolute()})")

    # Multi-Language Side-by-Side Markdown, PDF, and EPUB
    targets_suffix = "-".join(target_langs)
    dual_md_path = book_output_dir / f"{base_name}_dual_{source_lang}_{targets_suffix}.md"
    md_lines = [f"# 📖 {base_name} - Multi-Lingual Book Edition\n"]
    for s in translated_segments:
        orig = s.get("original_text", s.get("original_ar", s.get("text", ""))).strip()
        md_lines.append(f"### Segment {s['id']}\n")
        md_lines.append(f"**{source_lang.upper()} (Source)**: {orig}\n")
        for target in target_langs:
            trans = s.get(f"text_{target}", "").strip()
            md_lines.append(f"**{target.upper()}**: {trans}\n")
        md_lines.append("---\n")
    dual_md_content = "\n".join(md_lines)
    dual_md_content = enrich_markdown_arabic(dual_md_content)
    dual_md_path.write_text(dual_md_content, encoding="utf-8")
    print(f"       Multi-Lingual Book (.md): {dual_md_path.name}")
    generated_files.append(f"Multi-Lingual Book (.md): [`{dual_md_path.name}`](file://{dual_md_path.absolute()})")

    # Dual PDF & EPUB
    dual_pdf_path = book_output_dir / f"{base_name}_dual_{source_lang}_{targets_suffix}.pdf"
    if export_to_pdf(dual_md_content, dual_pdf_path, title=f"{base_name} Dual Edition"):
        generated_files.append(f"Multi-Lingual Book (.pdf): [`{dual_pdf_path.name}`](file://{dual_pdf_path.absolute()})")

    dual_epub_path = book_output_dir / f"{base_name}_dual_{source_lang}_{targets_suffix}.epub"
    if export_to_epub(dual_md_content, dual_epub_path, title=f"{base_name} Dual Edition"):
        generated_files.append(f"Multi-Lingual Book (.epub): [`{dual_epub_path.name}`](file://{dual_epub_path.absolute()})")

    # 6. Save Markdown Summary Report
    summary_report_path = book_output_dir / f"{base_name}_summary.md"
    
    files_list_md = "\n".join([f"{idx+1}. {item}" for idx, item in enumerate(generated_files)])
    report_md = f"""# 📚 Book Translation & Transcript Report

- **Source Book**: `{book_path.name}`
- **Source Format**: `{book_path.suffix.upper()}`
- **Processed At**: `{datetime.now(UTC).strftime('%Y-%m-%d %H:%M:%S')} UTC`
- **Total Paragraph Chunks**: `{len(translated_segments)}`
- **Total Raw Characters**: `{len(raw_text):,}`

---

## 📦 Generated Book Files

1. **📄 Extracted {source_lang.upper()} Text Transcript**: [`{extracted_src_path.name}`](file://{extracted_src_path.absolute()})
{files_list_md}

---

## 🤖 AI Engine Stack

- **Extraction**: `pypdf` / `python-docx` / UTF-8 Text Engine
- **Primary AI LLM**: `Qwen2.5-3B-Instruct-8bit` / Local Fallback MarianMT (MLX Metal/MPS GPU)
- **Translation Strategy**: Multi-Pass Agentic Self-Refinement (Draft -> Critique & Refine)
"""
    summary_report_path.write_text(report_md, encoding="utf-8")
    print(f"       Markdown Summary Report: {summary_report_path.name}")

    print(f"\n{'='*60}")
    print(f"  🎉 Book Translation Complete! Output files saved in '{book_output_dir}/'")
    print(f"{'='*60}\n")
    return book_output_dir
