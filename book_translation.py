"""
Book Translation Module (Arabic → German / English / Both)
==========================================================

Supports translating book files (.txt, .pdf, .docx) using the 3-Tier
fallback translation pipeline with batching, live JSON backup, and resume.
Outputs:
  - Translated Text / Markdown book files
  - Word (.docx) documents (if python-docx is installed)
  - Full structured JSON backup
"""

import json
from pathlib import Path
from typing import List, Dict, Any


def extract_text_from_book(file_path: Path) -> str:
    """Extract full raw text from PDF, DOCX, or TXT book file."""
    ext = file_path.suffix.lower()
    
    if ext == ".txt" or ext == ".md":
        return file_path.read_text(encoding="utf-8")
        
    elif ext == ".pdf":
        try:
            import pypdf
            reader = pypdf.PdfReader(str(file_path))
            text_parts = []
            for i, page in enumerate(reader.pages):
                p_text = page.extract_text()
                if p_text:
                    text_parts.append(p_text)
            return "\n\n".join(text_parts)
        except ImportError:
            raise RuntimeError("pypdf is required to read PDF books. Install with: uv add pypdf")
            
    elif ext == ".docx":
        try:
            import docx
            doc = docx.Document(str(file_path))
            return "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            raise RuntimeError("python-docx is required to read DOCX books. Install with: uv add python-docx")
            
    else:
        raise ValueError(f"Unsupported book format: {ext}. Supported formats: .txt, .pdf, .docx, .md")


def chunk_text_into_paragraphs(text: str, max_chunk_chars: int = 1000) -> List[Dict[str, Any]]:
    """Split book text into coherent paragraph segments for translation."""
    raw_paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    
    chunks = []
    current_chunk = []
    current_length = 0
    
    for p in raw_paragraphs:
        if current_length + len(p) > max_chunk_chars and current_chunk:
            combined = "\n".join(current_chunk)
            chunks.append({
                "id": len(chunks) + 1,
                "text": combined,
                "original_ar": combined,
            })
            current_chunk = [p]
            current_length = len(p)
        else:
            current_chunk.append(p)
            current_length += len(p)
            
    if current_chunk:
        combined = "\n".join(current_chunk)
        chunks.append({
            "id": len(chunks) + 1,
            "text": combined,
            "original_ar": combined,
        })
        
    return chunks


def translate_book_interactive(book_path: Path, output_dir: Path, translate_segments_fn, args) -> Path:
    """Interactively process and translate a book from any source language to multiple target languages with dedicated directory & live backup."""
    base_name = book_path.stem
    book_output_dir = output_dir / "books" / base_name
    book_output_dir.mkdir(parents=True, exist_ok=True)
    
    print(f"\n{'='*60}")
    print(f"  📚 Book Translation Mode: {book_path.name}")
    print(f"  📁 Dedicated Book Output Directory: {book_output_dir}")
    print(f"{'='*60}")
    
    source_lang = getattr(args, "source_lang", "ar").lower()
    target_langs_str = getattr(args, "target_lang", "de,en").lower()

    # Determine target languages
    if getattr(args, "url", None) is not None or getattr(args, "target_lang", None) is not None or getattr(args, "source_lang", None) is not None:
        target_langs = [l.strip() for l in target_langs_str.split(",") if l.strip()]
    else:
        # Prompt target selection
        print("\n  Select Target Translation Language(s):")
        print("    [1] German (Deutsch)")
        print("    [2] English")
        print("    [3] Both German & English (Dual Translation)")
        lang_choice = input("\n  Enter choice [1-3] (default: 3 Both): ").strip()
        if lang_choice == "1":
            target_langs = ["de"]
        elif lang_choice == "2":
            target_langs = ["en"]
        else:
            target_langs = ["de", "en"]
            
    print(f"  Selected translation: {source_lang.upper()} → {', '.join(t.upper() for t in target_langs)}")
    
    # 2. Extract & Chunk Text
    print(f"\n[1/3] Extracting text from {book_path.name}...")
    raw_text = extract_text_from_book(book_path)
    segments = chunk_text_into_paragraphs(raw_text)
    print(f"       Extracted {len(segments)} paragraph chunks from book.")
    
    # 3. Translate Segments via Pipeline with dedicated directory for live JSON backups
    print(f"\n[2/3] Translating book paragraphs using 3-Tier AI Engine...")
    translated_segments = translate_segments_fn(segments, book_output_dir, args, source_lang=source_lang, target_langs=target_langs)
    
    # 4. Save Extracted Raw Text Transcript File
    extracted_src_path = book_output_dir / f"{base_name}_extracted_{source_lang}.txt"
    extracted_src_path.write_text(raw_text, encoding="utf-8")
    print(f"       Extracted {source_lang.upper()} Transcript (.txt): {extracted_src_path.name}")
 
    # 5. Save Output Book Files
    print(f"\n[3/3] Saving translated book deliverables & summary report...")
    
    # Save each target language txt/docx
    generated_files = []
    for target in target_langs:
        txt_path = book_output_dir / f"{base_name}_translated_{target}.txt"
        t_text = "\n\n".join([s.get(f"text_{target}", s.get("text", "")) for s in translated_segments])
        txt_path.write_text(t_text, encoding="utf-8")
        print(f"       {target.upper()} Book (.txt): {txt_path.name}")
        generated_files.append(f"{target.upper()} Book (.txt): [`{txt_path.name}`](file://{txt_path.absolute()})")
        
        try:
            import docx
            doc = docx.Document()
            doc.add_heading(f"{base_name} - {target.upper()} Translation", 0)
            for s in translated_segments:
                p_text = s.get(f"text_{target}", s.get("text", "")).strip()
                if p_text:
                    doc.add_paragraph(p_text)
            docx_path = book_output_dir / f"{base_name}_translated_{target}.docx"
            doc.save(str(docx_path))
            print(f"       {target.upper()} Book (.docx): {docx_path.name}")
            generated_files.append(f"{target.upper()} Book (.docx): [`{docx_path.name}`](file://{docx_path.absolute()})")
        except Exception:
            pass

    # Multi-Language Side-by-Side Markdown
    targets_suffix = "-".join(target_langs)
    dual_md_path = book_output_dir / f"{base_name}_dual_{source_lang}_{targets_suffix}.md"
    md_lines = [f"# 📖 {base_name} - Multi-Lingual Book Edition\n"]
    for s in translated_segments:
        orig = s.get("original_text", s.get("original_ar", s.get("text", ""))).strip()
        md_lines.append(f"### Segment {s['id']}\n")
        md_lines.append(f"**{source_lang.upper()} (Source)**: {orig}\n")
        for target in target_langs:
            trans = s.get(f"text_{target}", "").strip()
            md_lines.append(f"**{target.upper()}**: {trans}\n")
        md_lines.append("---\n")
    dual_md_path.write_text("\n".join(md_lines), encoding="utf-8")
    print(f"       Multi-Lingual Book (.md): {dual_md_path.name}")
    generated_files.append(f"Multi-Lingual Book (.md): [`{dual_md_path.name}`](file://{dual_md_path.absolute()})")

    # 6. Save Markdown Summary Report
    summary_report_path = book_output_dir / f"{base_name}_summary.md"
    from datetime import datetime
    
    files_list_md = "\n".join([f"{idx+1}. {item}" for idx, item in enumerate(generated_files)])
    report_md = f"""# 📚 Book Translation & Transcript Report

- **Source Book**: `{book_path.name}`
- **Source Format**: `{book_path.suffix.upper()}`
- **Processed At**: `{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}`
- **Total Paragraph Chunks**: `{len(translated_segments)}`
- **Total Raw Characters**: `{len(raw_text):,}`

---

## 📦 Generated Book Files

1. **📄 Extracted {source_lang.upper()} Text Transcript**: [`{extracted_src_path.name}`](file://{extracted_src_path.absolute()})
{files_list_md}

---

## 🤖 AI Engine Stack

- **Extraction**: `pypdf` / `python-docx` / UTF-8 Text Engine
- **Primary AI LLM**: `Qwen2.5-3B-Instruct-8bit` / Local Fallback MarianMT (MLX Metal/MPS GPU)
- **Translation Strategy**: Multi-Pass Agentic Self-Refinement (Draft -> Critique & Refine)
"""
    summary_report_path.write_text(report_md, encoding="utf-8")
    print(f"       Markdown Summary Report: {summary_report_path.name}")

    print(f"\n{'='*60}")
    print(f"  🎉 Book Translation Complete! Output files saved in '{book_output_dir}/'")
    print(f"{'='*60}\n")
    return book_output_dir
